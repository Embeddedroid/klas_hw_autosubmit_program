from selenium import webdriver # 웹 브라우저 자동화 모듈
import time

# 웹 브라우저 자동화
driver = webdriver.Chrome()

driver.get('https://klas.kw.ac.kr/') # 사이트로 이동
# klas 로그인
kwid = '****'
kwpw = '*****'
driver.find_element_by_id('loginId').send_keys(kwid)
driver.find_element_by_id('loginPwd').send_keys(kwpw)
driver.find_element_by_class_name('btn').click()

# 주요 일정 출력
time.sleep(1)

driver.find_element_by_class_name('more_btn').click()

subjectable = driver.find_element_by_xpath("//*[@id='tableTest']/tbody")

for tr in subjectable.find_elements_by_tag_name("tr"):
    td = tr.find_elements_by_tag_name("td")
    s = "{} , {}, {}\n".format(td[0].text, td[1].text , td[2].text)
    print (s)

import imapclient # 메일 모듈
import pyzmail # 메일 모듈

# 메일에 있는 첨부 파일 저장하기
mail_tf = int(input('메일으로부터 다운로드 받을 pdf 첨부파일이 있나요? (yes : 1 입력): '))

if mail_tf==1:
    imap = imapclient.IMAPClient('imap.gmail.com', ssl=True)
    imap = imapclient.IMAPClient('imap.gmail.com')
    myid_gmail='******'
    mypw_gmail='@@@@@@@'
    imap.login(myid_gmail, mypw_gmail)  # id, password

    imap.select_folder('[Gmail]/전체보관함', readonly=False)
    s=input('과제 파일이 담긴 메일 제목을 입력하세요:')
    s=str(s)
    uids = imap.gmail_search('subject:'+s)  # '반도체소자1 hw3'란 제목을 포함한 메일 검색

    rmsgs = imap.fetch(uids, ['BODY[]'])
    message = pyzmail.PyzMessage.factory(rmsgs[uids[0]][b'BODY[]'])
    subject = message.get_subject()
    senders = message.get_addresses('from')
    print('제목:', message.get_subject())    # 제목 출력

    # 첨부 파일 저장
    
    mail_index = 0
    for part in message.mailparts:
        if part.filename:
            print('saving..', part.filename, part.type)
            cont = part.get_payload()
            if part.type.startswith('text/'):
                open(part.filename, 'w').write(cont)
            else:
                open(part.filename, 'wb').write(cont)
            mailheading = part.filename
            
# word 로 과제 표지 만들기 -> pdf 변환하기
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm
import datetime

heading = str(input("과제 제목 : "))
name = str(input("이름 : "))
studcode = str(input("학번 : "))

doc = Document()

p = doc.add_heading(heading, 0)

table = doc.add_table(rows=3, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
table.rows[0].cells[0].text = '성명'
table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True

table.rows[0].cells[1].text = name

table.rows[1].cells[0].text = '학번'
table.rows[1].cells[0].paragraphs[0].runs[0].font.bold = True

table.rows[1].cells[1].text = studcode

date = datetime.datetime.today()
table.rows[2].cells[0].text = '날짜'
table.rows[2].cells[0].paragraphs[0].runs[0].font.bold = True

table.rows[2].cells[1].text = str(date.year)+'.'+str(date.month)+'.'+str(date.day)



doc.save(heading+'.docx')

# 표지 pdf 변환 과정
import os
from win32com.client import Dispatch

wordapp = Dispatch("Word.Application")
wordapp.Visible = False

fpath = os.path.join(os.getcwd(), heading + ".docx")  # 파일 경로
myDoc = wordapp.Documents.Open(FileName=fpath)

pdf_path = os.path.join(os.getcwd(), heading + ".pdf")  # PDF로 저장
myDoc.SaveAs(pdf_path, FileFormat=17) # magic number

myDoc.Close()
wordapp.Quit()

# 제출할 이미지 파일 pdf로 만들기
image_tf = int(input("이미지 파일을 같이 제출하시겠습니까? (yes : 1 입력): "))


import glob # 파일 경로 이동
from reportlab.pdfgen import canvas # 빈 페이지
from reportlab.lib.pagesizes import A4 # 용지 크기
from PIL import Image

if(image_tf == 1):
    imageheading = str(input("이미지 pdf 제목 설정해주세요 : "))
    imagepath = str(input("이미지 파일들이 속한 경로 적어주세요 (복붙하시면 됩니다): "))
    imagepath = imagepath.replace('\\','/')
    # 이미지 파일 목록을 가져온다
    flist = glob.glob(imagepath + '/*.PNG')  # 이미지가 있는 폴더로 경로 변경 해주세요

    # 파일 저장을 위한 Canvas 객체를 만든다
    pgsize = A4  # 용지 설정
    c = canvas.Canvas(imageheading + '.pdf', pagesize=pgsize)  # 출력 파일, 페이지 크기 변경
    w2 = pgsize[0]      # 용지의 폭
    for fpath in flist:
        im = Image.open(fpath)  # 이미지를 읽는다
        w, h = im.size  # 가로, 세로 이미지 크기를 가져온다
        r = w / h       # 가로, 세로 비를 계산한다
        h2 = w2 / r     # 용지 폭에 맞추어 높이를 계산한다(비율 유지)
        c.drawImage(fpath, 0, 0, w2, h2)  # 이미지를 페이지에 그린다
        c.showPage()    # 한 페이지 정리하고 다음 페이지로 넘어간다
    c.save()

# pdf 병합하기
from PyPDF2 import PdfFileMerger, PdfFileReader

# image 파일이 있고, mail 파일도 있는 경우
if image_tf ==1: 
    pdfs = [heading + '.pdf', imageheading + '.pdf']
    
    if mail_tf == 1:
        pdfs.append(mailheading)
        
# image 파일만 있는 경우
elif mail_tf == 1: 
    pdfs=[heading + '.pdf']
    pdfs.append(mailheading)
    
# 둘다 없는 경우
else :
    pdfs[heading + '.pdf']
    
merger = PdfFileMerger()

for filename in pdfs:
    merger.append(PdfFileReader(open(filename, 'rb')))

merger.write(heading + ".pdf")

# 과제 제출 창 띄우기
driver.find_element_by_class_name('navbar-toggler.navtoggler').click() # 목록 버튼 클릭
time.sleep(1)
driver.find_elements_by_class_name('depth03ul')[3].find_elements_by_tag_name('a')[5].click() # 과제제출 버튼 클릭
i = int(input("\n0: 기초회로실험1\n1: 디지털공학\n2: 진로탐색및설계\n3: 전자기학\n4: 컴퓨터언어\n5: 회로이론\n6: 공학수학1\n7: 반도체소자1\n 제출할 과목을 선택하세요 : "))

driver.find_elements_by_class_name('form-control.form-control-sm')[1].find_elements_by_tag_name('option')[i].click() # 과제제출 버튼 클릭

reportable = driver.find_element_by_class_name('AType')
godfix=0
for tbody in reportable.find_elements_by_tag_name("tbody"):
    time.sleep(1)
    tr=tbody.find_element_by_tag_name("tr")
    time.sleep(1)
    td = tr.find_elements_by_tag_name("td")
    
    s = "{} , {}, {}\n".format(td[1].text , td[2].text, td[3].text)
    print(s)
    select = int(input('해당 항목으로 과제 제출하시겠습니까? : (0:no , 1:yes) '))
    if select == 1:
        if td[3].text=="제출" :
            fix = int(input('수정하시겠습니까? (1:yes, 0:no): '))
            if fix == 1:
                godfix=1
            else:
                break
        td[4].find_element_by_class_name('btn2.btn-gray').click()
        break



# 과제 파일 첨부하기
time.sleep(1)
import os
filepath = os.getcwd()+'\\'+heading+'.pdf'


# 과제 제목, 본문 입력
if godfix==1: # 수정하는 경우
    driver.find_element_by_xpath('//*[@id="appModule"]/div/div[5]/button[1]').click()
else: # 처음 제출하는 경우
    driver.find_element_by_css_selector("input[type='text']").send_keys(heading)
    driver.find_element_by_id('textarea').send_keys("안녕하세요.\n"+heading+' 제출합니다')

driver.find_element_by_css_selector("input[type='file']").send_keys(filepath)

if godfix==1: # 수정하는 경우
    driver.find_element_by_xpath('//*[@id="appModule"]/div/div[3]/button[1]').click() # 제출 버튼 클릭
else: # 처음 제출하는 경우
    driver.find_element_by_xpath('//*[@id="appModule"]/div[4]/button[1]').click() # 제출 버튼 클릭

